"""Document ingestion: turn a file into payloads the classifier can consume.

Each input file is converted into a list of payloads, where a payload is either:

- ``str``           -- extracted plain text (native-text inputs), or
- ``BinaryContent`` -- raw image bytes for the vision model to read (OCR is done
                       by the VLM, so there is no local OCR engine).

A single document may produce several payloads (e.g. one image per scanned page).
"""

from __future__ import annotations

from pathlib import Path

from pydantic_ai import BinaryContent

from .config import settings

Payload = str | BinaryContent

# Extensions handled as native text vs. as images.
_IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}


class IngestError(Exception):
    """Raised when a file cannot be ingested."""


def ingest(path: str | Path) -> list[Payload]:
    """Convert ``path`` into a list of text/image payloads for classification."""
    path = Path(path)
    if not path.is_file():
        raise IngestError(f"not a file: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _ingest_pdf(path)
    if suffix in _IMAGE_MEDIA_TYPES:
        return [_image_payload(path.read_bytes(), _IMAGE_MEDIA_TYPES[suffix])]
    if suffix == ".docx":
        return [_ingest_docx(path)]
    if suffix == ".doc":
        raise IngestError(
            ".doc (legacy Word) is not supported yet -- please convert to PDF or .docx"
        )
    raise IngestError(f"unsupported file type: {suffix!r} ({path.name})")


def _image_payload(data: bytes, media_type: str) -> BinaryContent:
    return BinaryContent(data=data, media_type=media_type)


def _ingest_docx(path: Path) -> str:
    """Extract text from a .docx (paragraphs + table cells)."""
    from docx import Document  # python-docx; imported lazily

    document = Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise IngestError(f"no text extracted from {path.name}")
    return text


def _ingest_pdf(path: Path) -> list[Payload]:
    """Ingest a PDF page by page.

    Text pages are extracted with pypdf. Pages with little/no extractable text
    are treated as scanned: the page is rasterized and emitted as an image
    payload for the VLM to read.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    payloads: list[Payload] = []
    scanned_page_indexes: list[int] = []

    for index, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if len(text) >= settings.pdf_text_min_chars:
            payloads.append(text)
        else:
            scanned_page_indexes.append(index)

    if scanned_page_indexes:
        payloads.extend(_rasterize_pdf_pages(path, scanned_page_indexes))

    if not payloads:
        raise IngestError(f"no content extracted from {path.name}")
    return payloads


def _rasterize_pdf_pages(path: Path, page_indexes: list[int]) -> list[BinaryContent]:
    """Rasterize the given (0-based) PDF pages to PNG image payloads.

    Requires Poppler on PATH (used by pdf2image). Imported lazily so the rest of
    the app works without Poppler when no scanned PDFs are involved.
    """
    try:
        from pdf2image import convert_from_path
    except Exception as exc:  # pragma: no cover - import guard
        raise IngestError(f"pdf2image is required for scanned PDFs: {exc}") from exc

    # Allow pointing at a Poppler bin dir when it isn't on PATH (see config).
    # pdf2image accepts poppler_path=None to mean "use PATH"; its type stub
    # omits the Optional, so silence the false positive.
    poppler_path: str | None = settings.poppler_path or None

    payloads: list[BinaryContent] = []
    for index in page_indexes:
        page_number = index + 1  # pdf2image is 1-based
        try:
            images = convert_from_path(
                str(path),
                first_page=page_number,
                last_page=page_number,
                dpi=settings.pdf_render_dpi,
                poppler_path=poppler_path,  # type: ignore[arg-type]
            )
        except Exception as exc:
            raise IngestError(
                "failed to rasterize scanned PDF page "
                f"{page_number} of {path.name}; is Poppler installed and on PATH? ({exc})"
            ) from exc
        for image in images:
            import io

            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            payloads.append(_image_payload(buffer.getvalue(), "image/png"))
    return payloads
